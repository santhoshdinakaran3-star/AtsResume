import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_optimized_resume(analysis_data: dict, extracted_data: dict) -> bytes:
    """Generate an ATS-optimized DOCX resume from analysis data."""
    doc = Document()
    
    # 1. Styles & Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Helper for adding headings
    def add_section_heading(text):
        h = doc.add_heading(text.upper(), level=1)
        h.style.font.size = Pt(12)
        h.style.font.bold = True
        h.style.font.name = 'Arial'

    # 0. Content optimization (already provided in analysis)
    improved_raw = analysis_data.get("improved_experience", [])
    
    # Map improved bullets back
    ai_experience = []
    if improved_raw:
        for exp in improved_raw:
            bullets = exp.get("improved_bullets", [])
            for b in bullets:
                ai_experience.append(b)
    else:
        ai_experience = analysis_data.get("extracted_experience", [])

    ai_skills = analysis_data.get("extracted_skills", [])
    ai_summary = analysis_data.get("ai_summary", "")

    # 2. Header
    full_text = extracted_data.get("full_text", "")
    pinfo = analysis_data.get("personal_info", {})
    name = pinfo.get("name") or full_text.split('\n')[0].strip()
    if len(name) > 40: name = "[YOUR NAME]"
    
    # Extract email and phone if not in pinfo
    email = pinfo.get("email")
    if not email:
        email_match = re.search(r"[\w\.\-]+@[\w\.\-]+\.\w+", full_text)
        email = email_match.group(0) if email_match else ""
        
    phone = pinfo.get("phone")
    if not phone:
        phone_match = re.search(r"[\+]?[\d\s\-\(\)]{10,15}", full_text)
        phone = phone_match.group(0) if phone_match else ""
        
    linkedin = pinfo.get("linkedin") or ""

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(name)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Arial'

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = []
    if email: contact_info.append(email)
    if phone: contact_info.append(phone)
    
    # Look for LinkedIn
    li = re.search(r"linkedin\.com/in/[a-zA-Z0-9\-]+", full_text, re.IGNORECASE)
    if li: contact_info.append(li.group(0))
    
    contact_para.add_run(" | ".join(contact_info))
    contact_para.style.font.size = Pt(10)
    contact_para.style.font.name = 'Arial'

    doc.add_paragraph() # Spacer

    # 3. Professional Summary
    add_section_heading("Professional Summary")
    summary_text = ai_summary if ai_summary else f"Result-driven professional with expertise in {', '.join(ai_skills[:5])}."
    
    p = doc.add_paragraph(summary_text)
    p.style.font.size = Pt(10)
    p.style.font.name = 'Arial'

    # 4. Experience
    add_section_heading("Experience")
    if ai_experience:
        for exp in ai_experience:
            p = doc.add_paragraph(exp, style='List Bullet')
            p.style.font.size = Pt(10)
            p.style.font.name = 'Arial'
    else:
        p = doc.add_paragraph("[Enter your relevant work experience here]", style='List Bullet')
    
    # 5. Skills
    add_section_heading("Technical Skills")
    if ai_skills:
        skills_para = doc.add_paragraph()
        run = skills_para.add_run(", ".join(ai_skills))
        run.font.size = Pt(10)
        run.font.name = 'Arial'
    
    # 6. Education
    add_section_heading("Education")
    edu_info = analysis_data.get("extracted_education", [])
    if edu_info:
        for edu in edu_info:
            p = doc.add_paragraph(edu, style='List Bullet')
            p.style.font.size = Pt(10)
            p.style.font.name = 'Arial'
    
    # Save to buffer
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream.getvalue()
